"""Read an attached file into the text of one context item.

This is the whole of the file boundary for Expert's transient context. Nothing above or
below it changes: a `ContextItem` is still a label and a body of text, still travels only
into a prompt, is still never chunked, never cited, and never checked against a quote.
What changed is where the text comes from — an attachment rather than a paste, because
nobody has the text of a report to hand, and everybody has the file.

**It knows nothing about what the document is.** No document type, no section names, no
expectation of headings, tables, or an author. It reads whatever the file says, in reading
order, and hands it over as prose. That is the reason a format with no declared structure
is safe here and refused for an upload: an uploaded document becomes citable blocks whose
meaning downstream depends on structure being read rather than inferred, and context
becomes a paragraph in a prompt.

Deliberately not in Chunker. Chunker's contract is declared structure in, citable blocks
out, and a flat string is neither. If a second tool needs this, that is the moment to move
it to `shared/` — not now.
"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

#: Formats a context attachment may carry.
#:
#: Wider than the analysis path on purpose, and the difference is not an oversight: an
#: uploaded document becomes citable blocks, so its structure has to be declared in the
#: file. Context is read once into a prompt, so a format that declares nothing loses
#: nothing. `.docx` is here as well as in the analysis path because a report or a set of
#: meeting notes is usually one, and refusing it would only send a reader to export the
#: same file as a PDF for an identical result.
CONTEXT_SUFFIXES = (".pdf", ".docx", ".txt", ".md")

CONTEXT_FORMAT_HINT = "PDF, DOCX, TXT, or MD"

#: Longest one attachment may be, in characters.
#:
#: A cap exists because attaching removed the friction that used to keep context small:
#: pasting a 200-page report was work, and dropping one in is not. Context is repeated in
#: every question's prompt — around eighty of them — so an unbounded attachment is an
#: unbounded run. Stated in characters rather than pages because two of the four formats
#: have no pages.
MAX_CONTEXT_CHARACTERS = 80_000

#: Longest all attachments may be together.
MAX_TOTAL_CONTEXT_CHARACTERS = 200_000


class ContextReadError(ValueError):
    """A file that cannot become context, with the reason a reader can act on."""


def read_context_text(file_path: str, *, filename: str | None = None) -> str:
    """The text of one attachment, page-delimited where the format has pages.

    Raises `ContextReadError` with a reason a reader can act on. It never returns an empty
    string: context that says nothing would sit in the prompt as a named source with
    nothing in it, and every question answered from it would be answered from silence.
    """
    name = Path(filename or file_path).name
    suffix = Path(name).suffix.lower()
    if suffix not in CONTEXT_SUFFIXES:
        raise ContextReadError(
            f"{name}: context may be {CONTEXT_FORMAT_HINT}."
        )

    if suffix == ".pdf":
        text = _read_pdf(file_path, name)
    elif suffix == ".docx":
        text = _read_docx(file_path, name)
    else:
        text = _read_plain(file_path, name)

    text = text.strip()
    if not text:
        raise ContextReadError(
            f"{name}: no text could be read from this file, so there would be nothing "
            "for a question to be answered from."
        )
    if len(text) > MAX_CONTEXT_CHARACTERS:
        raise ContextReadError(
            f"{name}: {len(text):,} characters of text, over the "
            f"{MAX_CONTEXT_CHARACTERS:,} limit for one context item. Context is read "
            "again for every question in the gate, so attach the part that bears on this "
            "review rather than the whole document."
        )
    return text


def total_context_length(texts: list[str]) -> int:
    """How much context a run carries, for the limit that covers all of it together."""
    return sum(len(text) for text in texts)


def _read_pdf(file_path: str, name: str) -> str:
    """Every page's text, in page order, each under a marker naming its page.

    PDFium, through `pypdfium2` — already this project's PDF engine, where it rasterizes
    converted slides. Text extraction is the same library asked a different question, so
    there is no second dependency and no service to call.

    The page marker is the only structure claimed, and it is the only one the format
    actually declares. Nothing here infers a heading, a table, or a column from where
    glyphs sit on the page: that inference is what makes a rendered format unsafe to
    treat as a document, and this text is not treated as one.
    """
    try:
        import pypdfium2 as pdfium
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ContextReadError(
            f"{name}: this deployment cannot read PDFs; attach DOCX, TXT, or MD instead."
        ) from exc

    try:
        document = pdfium.PdfDocument(file_path)
    except Exception as exc:  # pypdfium2 raises its own error types
        raise ContextReadError(
            f"{name}: this file could not be opened as a PDF."
        ) from exc

    pages: list[str] = []
    try:
        for number, page in enumerate(document, start=1):
            text_page = page.get_textpage()
            try:
                # `get_text_bounded` on current pypdfium2, `get_text_range` before it.
                # Both return the page's text; asked in this order so either works.
                reader = getattr(text_page, "get_text_bounded", None) or getattr(
                    text_page, "get_text_range"
                )
                page_text = (reader() or "").strip()
            finally:
                # Closed as we go rather than left to the finalizer. This runs inside a
                # long-lived API process, and a review may attach several hundred pages.
                text_page.close()
                page.close()
            if page_text:
                pages.append(f"## Page {number}\n\n{page_text}")
    except Exception as exc:
        raise ContextReadError(f"{name}: this PDF could not be read.") from exc
    finally:
        document.close()

    if not pages:
        # The common case is a scan: pages of images with no text layer. Said plainly,
        # because the alternative is a context item that exists and answers nothing.
        raise ContextReadError(
            f"{name}: this PDF has no text layer, which usually means it is a scan or a "
            "set of page images. Attach a text-based PDF, or the original document."
        )
    return "\n\n".join(pages)


def _read_docx(file_path: str, name: str) -> str:
    """A DOCX's paragraphs and table cells as flat text, in document order.

    Deliberately not Chunker's parser. That one returns citable blocks with declared
    structure, which is the right answer for a document under analysis and the wrong shape
    for a prompt paragraph. Reading it flat here keeps the two paths from being confused
    for one another: nothing produced by this function can be cited.
    """
    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise ContextReadError(
            f"{name}: this deployment cannot read DOCX files."
        ) from exc

    try:
        document = Document(file_path)
    except Exception as exc:
        raise ContextReadError(
            f"{name}: this file could not be opened as a DOCX."
        ) from exc

    parts = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                # Pipes because a row read as one sentence merges its columns, and a
                # reader of the prompt should see where one cell ends.
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _read_plain(file_path: str, name: str) -> str:
    """A text or markdown file as it was written.

    Markdown is passed through untouched: it is already text, and its structure is the
    author's own rather than something inferred here.
    """
    try:
        return Path(file_path).read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        raise ContextReadError(f"{name}: this file could not be read.") from exc
