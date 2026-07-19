from __future__ import annotations

import base64
import hashlib
import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document

from services.chunker import blocks_to_dicts, run_pipeline


PNG_1X1 = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


class ChunkerImageTests(unittest.TestCase):
    def test_docx_raster_survives_as_a_portable_image_block(self) -> None:
        document = Document()
        document.add_paragraph("Figure context")
        document.add_picture(BytesIO(PNG_1X1))

        with tempfile.TemporaryDirectory() as directory:
            path = Path(directory) / "visual.docx"
            document.save(path)
            blocks = run_pipeline(str(path), "visual")

        image_block = next(block for block in blocks if block.block_type == "image")
        self.assertEqual(image_block.content, "[image]")
        self.assertIsNotNone(image_block.image)
        assert image_block.image is not None
        decoded = base64.b64decode(image_block.image.data_base64)
        self.assertEqual(decoded, PNG_1X1)
        self.assertEqual(image_block.image.media_type, "image/png")
        self.assertEqual(image_block.image.source_media_type, "image/png")
        self.assertEqual(image_block.image.sha256, hashlib.sha256(PNG_1X1).hexdigest())
        self.assertNotIn("image_rel_id", image_block.structural_meta)
        self.assertEqual(
            blocks_to_dicts([image_block])[0]["image"]["data_base64"],
            image_block.image.data_base64,
        )


if __name__ == "__main__":
    unittest.main()
