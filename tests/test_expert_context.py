"""Any file in, one body of prose out, and it knows nothing about the document.

The reader is the whole file boundary for Expert's transient context. What it must not
become is a parser with opinions: no document type, no section names, no expectation of
headings or tables. It reads what the file says and hands it over, which is exactly why a
format with no declared structure is safe here and refused for an upload.
"""

from __future__ import annotations

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from services.expert import (
    CONTEXT_SUFFIXES,
    MAX_CONTEXT_CHARACTERS,
    ContextReadError,
    read_context_text,
    total_context_length,
)


def write_pdf(path: Path, pages: list[list[str]]) -> None:
    """A real PDF, written by hand so the test needs no extra dependency."""
    objs = ["<< /Type /Catalog /Pages 2 0 R >>"]
    kids = " ".join(f"{4 + index * 2} 0 R" for index in range(len(pages)))
    objs.append(f"<< /Type /Pages /Kids [{kids}] /Count {len(pages)} >>")
    objs.append("<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>")
    for index, lines in enumerate(pages):
        text = (
            "BT /F1 12 Tf 72 720 Td 16 TL\n"
            + "\n".join(f"({line}) Tj T*" for line in lines)
            + "\nET"
        )
        objs.append(
            "<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            f"/Resources << /Font << /F1 3 0 R >> >> /Contents {5 + index * 2} 0 R >>"
        )
        objs.append(f"<< /Length {len(text)} >>\nstream\n{text}\nendstream")
    body, offsets = "%PDF-1.4\n", []
    for number, obj in enumerate(objs, start=1):
        offsets.append(len(body))
        body += f"{number} 0 obj\n{obj}\nendobj\n"
    start = len(body)
    body += f"xref\n0 {len(objs) + 1}\n0000000000 65535 f \n"
    for offset in offsets:
        body += f"{offset:010d} 00000 n \n"
    body += (
        f"trailer\n<< /Size {len(objs) + 1} /Root 1 0 R >>\n"
        f"startxref\n{start}\n%%EOF\n"
    )
    path.write_bytes(body.encode("latin-1"))


class ContextReaderTests(unittest.TestCase):
    def setUp(self) -> None:
        self.directory = TemporaryDirectory()
        self.addCleanup(self.directory.cleanup)
        self.root = Path(self.directory.name)

    def test_a_pdf_becomes_page_delimited_text(self) -> None:
        """The page marker is the only structure claimed, and the only one declared."""
        path = self.root / "report.pdf"
        write_pdf(path, [
            ["Target Product Profile", "Dosing regimen: one dose, annually."],
            ["Manufacturing", "Cost of goods below USD 2.00 per dose."],
        ])
        text = read_context_text(str(path))
        self.assertIn("## Page 1", text)
        self.assertIn("## Page 2", text)
        self.assertIn("one dose, annually", text)
        self.assertLess(text.index("## Page 1"), text.index("## Page 2"))

    def test_nothing_is_inferred_beyond_the_page(self) -> None:
        """No heading, no table, no column — inferring those is what makes a rendered
        format unsafe to treat as a document, and this text is not treated as one."""
        path = self.root / "report.pdf"
        write_pdf(path, [["SECTION ONE", "Shelf life 36 months", "Storage 2-8 C"]])
        text = read_context_text(str(path))
        # The only markdown syntax present is the page marker this reader added.
        self.assertEqual(text.count("#"), 2)
        self.assertNotIn("|", text)
        self.assertNotIn("**", text)

    def test_a_pdf_with_no_text_layer_fails_loudly(self) -> None:
        """A scan. Silence here would be a named source that answers nothing."""
        path = self.root / "scan.pdf"
        write_pdf(path, [[]])
        with self.assertRaises(ContextReadError) as caught:
            read_context_text(str(path))
        self.assertIn("no text layer", str(caught.exception))

    def test_a_file_that_is_not_a_pdf_fails_with_its_name(self) -> None:
        path = self.root / "broken.pdf"
        path.write_bytes(b"not a pdf at all")
        with self.assertRaises(ContextReadError) as caught:
            read_context_text(str(path))
        self.assertIn("broken.pdf", str(caught.exception))

    def test_a_docx_becomes_flat_text_with_its_table_rows(self) -> None:
        """Read flat on purpose: nothing this returns can be cited, unlike Chunker's
        blocks, so the two paths cannot be mistaken for one another."""
        path = self.root / "summary.docx"
        document = Document()
        document.add_paragraph("Stability summary")
        table = document.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "Shelf life"
        table.rows[0].cells[1].text = "24 months"
        document.save(str(path))

        text = read_context_text(str(path))
        self.assertIn("Stability summary", text)
        # Pipes so a row does not read as one merged sentence.
        self.assertIn("Shelf life | 24 months", text)

    def test_markdown_and_text_pass_through_untouched(self) -> None:
        for name, body in (("notes.md", "# Notes\n\nAgreed 24 months."), ("notes.txt", "Plain.")):
            path = self.root / name
            path.write_text(body, encoding="utf-8")
            self.assertEqual(read_context_text(str(path)), body)

    def test_an_unsupported_format_is_refused_by_name(self) -> None:
        path = self.root / "deck.pptx"
        path.write_text("x", encoding="utf-8")
        with self.assertRaises(ContextReadError) as caught:
            read_context_text(str(path))
        self.assertIn("deck.pptx", str(caught.exception))
        self.assertIn("PDF, DOCX, TXT, or MD", str(caught.exception))

    def test_an_empty_file_is_refused(self) -> None:
        path = self.root / "empty.txt"
        path.write_text("   \n", encoding="utf-8")
        with self.assertRaises(ContextReadError):
            read_context_text(str(path))

    def test_an_oversized_attachment_is_refused_with_the_limit(self) -> None:
        """Attaching removed the friction that kept pasted context small, and context is
        read again for every question in the gate."""
        path = self.root / "huge.txt"
        path.write_text("x" * (MAX_CONTEXT_CHARACTERS + 1), encoding="utf-8")
        with self.assertRaises(ContextReadError) as caught:
            read_context_text(str(path))
        self.assertIn(f"{MAX_CONTEXT_CHARACTERS:,}", str(caught.exception))

    def test_the_name_in_an_error_is_the_reader_facing_one(self) -> None:
        """The route reads from a temp file, whose name means nothing to a user."""
        path = self.root / "tmpabc123.pdf"
        path.write_bytes(b"not a pdf")
        with self.assertRaises(ContextReadError) as caught:
            read_context_text(str(path), filename="CMC Report.pdf")
        self.assertIn("CMC Report.pdf", str(caught.exception))
        self.assertNotIn("tmpabc123", str(caught.exception))

    def test_the_total_is_measured_across_attachments(self) -> None:
        self.assertEqual(total_context_length(["abc", "de"]), 5)

    def test_the_supported_set_is_wider_than_the_analysis_path(self) -> None:
        """Deliberate: context is never chunked, cited, or quote-checked, so a format
        that declares no structure loses nothing here."""
        from services.chunker import DOCUMENT_SUFFIXES

        self.assertIn(".pdf", CONTEXT_SUFFIXES)
        self.assertNotIn(".pdf", DOCUMENT_SUFFIXES)
        self.assertNotIn(".pptx", CONTEXT_SUFFIXES)


class WebParityTests(unittest.TestCase):
    def test_the_browser_offers_exactly_what_the_service_reads(self) -> None:
        """Two lists, one set. A picker advertising a format the service refuses is an
        error a user only meets after choosing a file."""
        source = (
            Path(__file__).resolve().parents[1] / "web" / "lib" / "document-formats.ts"
        ).read_text(encoding="utf-8")
        declared = source.split("CONTEXT_SUFFIXES = [")[1].split("]")[0]
        for suffix in CONTEXT_SUFFIXES:
            self.assertIn(f'"{suffix}"', declared, suffix)
        self.assertEqual(declared.count('"'), len(CONTEXT_SUFFIXES) * 2)


if __name__ == "__main__":
    unittest.main()
